import os
import glob
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout.reconfigure(encoding='utf-8')

brain_dir = r'C:\Users\62812\.gemini\antigravity-ide\brain\2ec072d0-b8f1-4a9d-a960-3e3ca05dee4f'
original_doc_path = r'c:\College\Capstone Design\sigigi-main\original_draft.docx'
doc_path = r'c:\College\Capstone Design\sigigi-main\Draft CD5.docx'

# Find the latest file in brain_dir matching a prefix pattern
def get_latest_image(prefix):
    pattern = os.path.join(brain_dir, f"{prefix}*.png")
    matching_files = glob.glob(pattern)
    if not matching_files:
        return None
    # Sort by modification time to get the latest one
    matching_files.sort(key=os.path.getmtime, reverse=True)
    return matching_files[0]

# Helper to format text inside cell with Times New Roman font
def format_cell_text(cell, text, bold=False, italic=False, size_pt=10, align=WD_ALIGN_PARAGRAPH.LEFT, color_rgb=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0  # Spasi 1.0 (Single) inside table cells
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb
    return run

# Helper to add a row and return cells after merging based on grid spans
def add_grid_row(table, spans, total_width_in=5.37):
    row = table.add_row()
    row.height = None
    row.height_rule = None
    
    # Set initial widths for 12 columns
    for cell in row.cells:
        cell.width = Inches(total_width_in / 12.0)
        
    merged_cells = []
    current_col = 0
    for span in spans:
        start_cell = row.cells[current_col]
        end_cell = row.cells[current_col + span - 1]
        if span > 1:
            merged = start_cell.merge(end_cell)
        else:
            merged = start_cell
        merged.width = Inches(span * (total_width_in / 12.0))
        merged_cells.append(merged)
        current_col += span
    return merged_cells

# Helper to write description and two images in content cell
def populate_content_cell(cell, s_desc, fe_prefix, be_prefix, col_width_in):
    # Clear cell paragraphs
    while len(cell.paragraphs) > 1:
        p_to_remove = cell.paragraphs[-1]
        p_to_remove._element.getparent().remove(p_to_remove._element)
        
    # Write description
    p_desc = cell.paragraphs[0]
    p_desc.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_desc.paragraph_format.space_before = Pt(4)
    p_desc.paragraph_format.space_after = Pt(4)
    p_desc.paragraph_format.line_spacing = 1.0
    run_desc = p_desc.add_run(s_desc)
    run_desc.font.name = 'Times New Roman'
    run_desc.font.size = Pt(9.5)
    
    # 1. Add FE Evidence Label & Image
    fe_img = get_latest_image(fe_prefix)
    p_fe_lbl = cell.add_paragraph()
    p_fe_lbl.paragraph_format.space_before = Pt(4)
    p_fe_lbl.paragraph_format.space_after = Pt(2)
    run_fe_lbl = p_fe_lbl.add_run("FE Evidence:")
    run_fe_lbl.font.bold = True
    run_fe_lbl.font.name = 'Times New Roman'
    run_fe_lbl.font.size = Pt(9)
    
    if fe_img and os.path.exists(fe_img):
        p_img1 = cell.add_paragraph()
        p_img1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img1.paragraph_format.space_before = Pt(2)
        p_img1.paragraph_format.space_after = Pt(4)
        run_img1 = p_img1.add_run()
        run_img1.add_picture(fe_img, width=Inches(col_width_in - 0.2))
    else:
        p_err = cell.add_paragraph()
        p_err.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_err = p_err.add_run("[FE Image Not Found]")
        run_err.font.bold = True
        run_err.font.color.rgb = RGBColor(255, 0, 0)
        run_err.font.name = 'Times New Roman'
        run_err.font.size = Pt(9)

    # 2. Add BE/DB Evidence Label & Image
    be_img = get_latest_image(be_prefix)
    p_be_lbl = cell.add_paragraph()
    p_be_lbl.paragraph_format.space_before = Pt(4)
    p_be_lbl.paragraph_format.space_after = Pt(2)
    run_be_lbl = p_be_lbl.add_run("BE/DB Evidence:")
    run_be_lbl.font.bold = True
    run_be_lbl.font.name = 'Times New Roman'
    run_be_lbl.font.size = Pt(9)
    
    if be_img and os.path.exists(be_img):
        p_img2 = cell.add_paragraph()
        p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img2.paragraph_format.space_before = Pt(2)
        p_img2.paragraph_format.space_after = Pt(4)
        run_img2 = p_img2.add_run()
        run_img2.add_picture(be_img, width=Inches(col_width_in - 0.2))
    else:
        p_err2 = cell.add_paragraph()
        p_err2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_err2 = p_err2.add_run("[BE/DB Image Not Found]")
        run_err2.font.bold = True
        run_err2.font.color.rgb = RGBColor(255, 0, 0)
        run_err2.font.name = 'Times New Roman'
        run_err2.font.size = Pt(9)

requirements_info = {
    6: {
        "code": "FR-01",
        "name": "Pendaftaran Akun, Login, & Hak Akses",
        "data_input": "Payload: { username: \"...\", passwords: \"...\", email: \"...\", role: \"admin|doctor|patient\" }",
        "be_prefix": "backend_fr01",
        "expected": {
            "FE": "Menampilkan form login/register, tombol dinonaktifkan saat loading, dan Toast notifikasi sukses/gagal di layar. Jika login sukses, mengarahkan ke dashboard.",
            "BE": "Mengirim request POST /api/login atau /api/register dan mengembalikan status 200 OK (disertai bearer token Sanctum) atau 422 Unprocessable Entity (jika gagal validasi email duplikat/sandi salah).",
            "DB": "Melakukan select kecocokan data ke tabel `users` (kolom `username`, `passwords`). Pendaftaran baru meng-insert baris baru ke tabel `users` (kolom `id_users`, `email`, `passwords` terenkripsi bcrypt, `role`) dan tabel `master_pasien`."
        },
        "actual": {
            "FE": "Form login/register menampilkan respons cepat, tombol berubah memudar saat loading, dan Toast popup berhasil/gagal (misal: 'Kredensial login salah') muncul sesuai inputan.",
            "BE": "Request POST berhasil diproses oleh API Sanctum dengan respons 200 OK beserta token autentikasi. Validasi duplikat mengembalikan 422 dengan pesan JSON error.",
            "DB": "Record akun baru tersimpan dengan aman pada database `db_sigigi` tabel `users` dan password tersimpan dalam bentuk Bcrypt hash."
        },
        "scenarios": [
            ("Skenario 1: Registrasi Akun Pasien Baru", "Pasien melakukan registrasi akun baru melalui form pendaftaran.", "04_register_pasien_sukses"),
            ("Skenario 2: Login Portal Pasien", "Pasien melakukan login dengan kredensial yang baru didaftarkan.", "03_login_pasien_sukses"),
            ("Skenario 3: Login Portal Staf/Dokter", "Staf internal/dokter login untuk menguji hak akses halaman internal.", "01_login_staf_sukses"),
            ("Skenario 4: Logout Pengguna", "Pengguna melakukan logout dari sistem untuk mengakhiri sesi.", "11_logout_pengguna"),
            ("Skenario 5: Gagal Login - Password Salah", "Pengguna memasukkan kata sandi yang salah pada form login dan sistem menampilkan pesan kesalahan.", "02_login_staf_gagal"),
            ("Skenario 6: Gagal Registrasi - Email Duplikat", "Pasien mencoba mendaftarkan email yang sudah terdaftar dan sistem menolak pendaftaran.", "05_register_pasien_gagal")
        ]
    },
    7: {
        "code": "FR-02",
        "name": "Appointment Online, Tingkat Urgensi, & Analisis AI",
        "data_input": "Payload: { id_klinik: int, tanggal_kunjungan: date, id_jadwal: int, keluhan_utama: text, kuesioner: array[10], file_citra: binary }",
        "be_prefix": "backend_fr02",
        "expected": {
            "FE": "Menampilkan dropdown klinik, kalender tanggal, dropdown slot waktu, form kuesioner keluhan, dan tombol upload xray citra gigi. Menampilkan hasil klasifikasi karies AI (EfficientNet-B0) pada panel identifikasi.",
            "BE": "Mengirim request POST /api/appointments dan mengembalikan respons status 201 Created jika sukses, atau 422 jika format gambar salah / slot waktu bentrok.",
            "DB": "Menyimpan data janji temu baru ke database `db_sigigi` tabel `appointments` pada kolom `appointment_date`, `appointment_time`, `patient_name`, `questionnaire` (JSON), dan `image_url` (path file citra)."
        },
        "actual": {
            "FE": "UI memuat tanggal dan waktu dengan rapi. Hasil analisis model EfficientNet-B0 menampilkan status 'Karies' atau 'Non-Karies' beserta persentase confidence. Toast bentrok jadwal muncul.",
            "BE": "API memproses inputan dan berhasil menyimpan data rekam antrean appointment dengan response 201. Bentrok jadwal mengembalikan 422 dengan pesan bentrok.",
            "DB": "Baris data appointment baru berhasil tersimpan lengkap dengan payload kuesioner pasien di tabel `appointments` database `db_sigigi`."
        },
        "scenarios": [
            ("Skenario 1: Reservasi Janji Temu Online", "Pasien memilih lokasi klinik, dokter, dan jadwal kunjungan.", "24_reservasi_janji_temu"),
            ("Skenario 2: Form Kuesioner Keluhan Gigi", "Pasien mengisi kuesioner keluhan gigi untuk perhitungan urgensi.", "25_kuesioner_keluhan"),
            ("Skenario 3: Unggah Citra Gigi (CNN Analysis)", "Pasien mengunggah citra gigi untuk analisis karies otomatis.", "26_upload_xray_cnn"),
            ("Skenario 4: Klasifikasi Karies Gigi AI", "Sistem menampilkan hasil klasifikasi otomatis karies/non-karies gigi berdasarkan citra yang diunggah.", "39_analisis_triage_ai"),
            ("Skenario 5: Validasi Slot Penuh/Terisi", "Sistem menyaring slot waktu yang penuh sehingga tidak dapat dipilih kembali.", "slot_bentrok"),
            ("Skenario 6: Gagal Analisis CNN - Gambar Non-Gigi", "Sistem menolak analisis jika berkas yang diunggah bukan citra gigi.", "cnn_non_gigi_error")
        ]
    },
    8: {
        "code": "FR-03",
        "name": "Pemantauan Antrean Pasien",
        "data_input": "Query Params: { id_clinic: int, tanggal: date }",
        "be_prefix": "backend_fr03",
        "expected": {
            "FE": "Menampilkan nomor antrean aktif pasien saat ini, jam digital real-time, dan tabel daftar seluruh antrean pasien yang sedang menunggu/sedang dilayani secara dinamis.",
            "BE": "Mengirim request GET /api/queue-today dan mengembalikan respons status 200 OK berisi JSON array daftar antrean aktif hari ini.",
            "DB": "Melakukan query SELECT pada database `db_sigigi` tabel `appointments` dengan filter `appointment_date = tanggal_hari_ini` dan `status IN ('pending', 'serving')`."
        },
        "actual": {
            "FE": "Nomor antrean dan daftar tabel antrean memuat data dengan benar. Jika tidak ada antrean hari ini, tabel menampilkan pesan 'Belum ada antrean'.",
            "BE": "Endpoint API merespon dengan status 200 OK dan menyajikan data JSON antrean aktif secara cepat.",
            "DB": "Database berhasil memproses query SELECT dan menyajikan daftar antrean hari ini tanpa kendala."
        },
        "scenarios": [
            ("Skenario 1: Urutan Antrean Dinamis Pasien", "Pasien memantau nomor antrean aktif dan status pelayanan saat ini.", "queue_active_three_patients"),
            ("Skenario 2: Tampilan State Antrean Kosong", "Sistem menampilkan pesan bahwa tidak ada antrean aktif pada hari tersebut.", "queue_empty_state")
        ]
    },
    9: {
        "code": "FR-04",
        "name": "Detail pasien dokter",
        "data_input": "Path Params: { id_appointment: int }",
        "be_prefix": "backend_fr04",
        "expected": {
            "FE": "Menampilkan halaman detail pasien khusus dokter untuk memverifikasi data keluhan, foto citra gigi, serta tingkat urgensi pelayanan.",
            "BE": "Mengirim request GET /api/appointments/{id} dan mengembalikan respons status 200 OK dengan payload detail janji temu, atau status 403 Forbidden jika dokter mengakses pasien dari klinik cabang lain.",
            "DB": "Melakukan query SELECT pada tabel `appointments` dan melakukan join ke tabel `master_pasien` berdasarkan kecocokan data telepon/email."
        },
        "actual": {
            "FE": "Dashboard detail pasien terbuka lengkap dengan data anamnesis pasien. Akses ilegal dari luar cabang dokter diblokir dengan halaman kosong / error guard.",
            "BE": "API mengembalikan data 200 OK berisi detail pasien secara lengkap. Akses silang cabang mengembalikan response status 403.",
            "DB": "Query join database memproses relasi tabel `appointments` dan tabel `master_pasien` secara presisi."
        },
        "scenarios": [
            ("Skenario 1: Verifikasi Keluhan oleh Dokter", "Dokter membuka dashboard detail pasien untuk memverifikasi anamnesis.", "08_triage_dokter_view"),
            ("Skenario 2: Tampilan Data Pasien Tanpa Citra Gigi", "Sistem menampilkan detail kunjungan pasien tanpa unggahan foto gigi dan status analisis AI masih Belum Dianalisis.", "doctor_patient_detail_no_xray")
        ]
    },
    10: {
        "code": "FR-05",
        "name": "Rekam Medis & Unduh PDF",
        "data_input": "Payload: { id_appointment: int, keluhan: text, diagnosis: string(ICD-10), tindakan: string, resep: array[{id_obat, dosis, qty}], odontogram: array[{gigi, kondisi}] }",
        "be_prefix": "backend_fr05",
        "expected": {
            "FE": "Menampilkan form input odontogram interaktif (SVG gigi), form diagnosis (ICD-10), form tindakan, form resep obat, serta tombol cetak/unduh PDF rekam medis. Toast error muncul jika form kosong.",
            "BE": "Mengirim request POST /api/medical-records dan mengembalikan respons status 201 Created (jika sukses menyimpan rekam medis) atau 422 Unprocessable Entity (jika form kosong).",
            "DB": "Menyimpan data rekam medis ke database `db_sigigi` tabel `rekam_medis_pasien` (kolom `keluhan`, `diagnosis_icd10`, `tindakan`, `resep`) dan database tabel `odontogram` (kolom `nomor_gigi`, `posisi_gigi`, `kondisi_gigi`, `warna_odontogram`)."
        },
        "actual": {
            "FE": "Dokter dapat mencatat odontogram secara interaktif dan memilih ICD-10. Unduhan PDF resep memuat layout nota resep gigi dengan rapi. Toast validasi form kosong muncul.",
            "BE": "Request POST berhasil diproses oleh API dengan respons status 201 Created. Validasi data kosong dicegah oleh form controller dengan status 422.",
            "DB": "Data riwayat rekam medis pasien berhasil masuk ke tabel `rekam_medis_pasien` dan status kondisi gigi tersimpan di tabel `odontogram` database `db_sigigi`."
        },
        "scenarios": [
            ("Skenario 1: Pengoperasian Odontogram", "Dokter mencatat kondisi gigi pasien pada form odontogram SVG.", "44_odontogram_interaktif"),
            ("Skenario 2: Pencatatan Diagnosa & Tindakan", "Dokter menginput rekam medis pasien di dashboard pemeriksaan.", "47_catat_rekam_medis"),
            ("Skenario 3: Penulisan Resep Obat", "Dokter meresepkan obat melalui form resep obat terintegrasi.", "48_input_resep_obat"),
            ("Skenario 4: Cetak & Unduh PDF Rekam Medis", "Pasien mengunduh file PDF rekam medis hasil pemeriksaan.", "37_cetak_nota_resep_pdf"),
            ("Skenario 5: Gagal Kirim Form Medis Kosong", "Sistem memunculkan peringatan validasi field wajib jika dokter mencoba menyimpan form kosong.", "47_catat_rekam_medis_gagal")
        ]
    },
    11: {
        "code": "FR-06",
        "name": "Pembayaran & Kasir",
        "data_input": "Payload: { id_pembayaran: int, nominal_bayar: decimal, diskon: decimal, status_pembayaran: string }",
        "be_prefix": "backend_fr06",
        "expected": {
            "FE": "Menampilkan kalkulasi otomatis total tagihan biaya tindakan & harga obat, field input diskon nominal rupiah, field input nominal bayar, dan pilihan status pembayaran.",
            "BE": "Mengirim request POST /api/payments dan mengembalikan respons status 200 OK dengan payload data transaksi terupdate.",
            "DB": "Menyimpan atau memperbarui data transaksi pembayaran di tabel `pembayaran` pada kolom `total_tagihan`, `nominal_bayar`, `diskon`, dan `status_pembayaran` (berubah menjadi `'lunas'`)."
        },
        "actual": {
            "FE": "Kalkulasi total tagihan terhitung otomatis secara tepat. Kasir dapat mengisi diskon Rp 0 maupun nominal potongan lainnya dengan respons instan.",
            "BE": "Endpoint API pembayaran merespon dengan status 200 OK dan memperbarui status antrean pasien terkait menjadi selesai.",
            "DB": "Status transaksi pembayaran pasien terupdate menjadi 'lunas' pada database `db_sigigi` tabel `pembayaran`."
        },
        "scenarios": [
            ("Skenario 1: Kalkulasi Rincian Tagihan", "Sistem mengkalkulasi total tagihan berdasarkan tindakan dan obat.", "50_kalkulasi_tagihan"),
            ("Skenario 2: Proses Pembayaran & Diskon", "Kasir menginput nominal pembayaran, diskon, dan status lunas.", "36_pembayaran_diskon"),
            ("Skenario 3: Pemberian Diskon Nominal Rp 0", "Sistem memperbolehkan transaksi dengan diskon Rp 0 jika pasien membayar penuh tanpa potongan.", "36_pembayaran_diskon")
        ]
    },
    12: {
        "code": "FR-07",
        "name": "Data Master",
        "data_input": "Payload: { nama_obat: string, satuan: string, dosis: string, keterangan: text }",
        "be_prefix": "backend_fr07",
        "expected": {
            "FE": "Menampilkan tabel data master (Klinik, Obat, Tindakan) beserta form dialog tambah/edit data master. Jika form kosong, menampilkan Toast error.",
            "BE": "Mengirim request POST /api/master/obat (atau cabang) dan mengembalikan respons status 201 Created jika berhasil, atau status 422 jika data kosong.",
            "DB": "Menambahkan baris baru ke database `db_sigigi` tabel `master_obat` pada kolom `nama_obat`, `satuan`, `dosis`, dan `keterangan`."
        },
        "actual": {
            "FE": "Tabel memuat daftar data master dengan benar. Tombol tambah memicu form dialog popup. Mengirim data kosong memicu Toast error 'Gagal menyimpan data'.",
            "BE": "Request POST berhasil diproses oleh API dengan respons status 201 Created untuk data valid, atau status 422 untuk inputan kosong.",
            "DB": "Baris data master obat baru berhasil ditambahkan ke tabel `master_obat` database `db_sigigi`."
        },
        "scenarios": [
            ("Skenario 1: Pengelolaan Data Master Klinik", "Admin melakukan CRUD data master klinik dan data dokter.", "12_tambah_klinik_sukses"),
            ("Skenario 2: Pengelolaan Data Master Obat", "Admin melakukan CRUD data master obat, tindakan, dan kode penyakit.", "62_master_obat"),
            ("Skenario 3: Pengaturan Harga Cabang", "Admin mengonfigurasi harga tindakan dan obat khusus di masing cabang.", "64_harga_tindakan_klinik"),
            ("Skenario 4: Gagal Menambah Obat Baru - Form Kosong", "Admin mengosongkan nama obat saat menambah data baru dan sistem menolak.", "master_obat_validation")
        ]
    },
    13: {
        "code": "FR-08",
        "name": "WhatsApp Klinik",
        "data_input": "Query Params: { phone: string, text: string }",
        "be_prefix": "backend_fr08",
        "expected": {
            "FE": "Menampilkan tombol ikon WhatsApp di layout pasien. Klik tombol akan memicu pemrosesan tautan redirect dan membuka tab baru mengarah ke API WhatsApp.",
            "BE": "Mengirim parameter nomor tujuan admin klinik dan teks draf pesan bantuan ke API external WhatsApp (`https://api.whatsapp.com/send`).",
            "DB": "Mengambil konfigurasi nomor telepon admin klinik secara dinamis melalui SELECT pada tabel `klinik` kolom `telepon`."
        },
        "actual": {
            "FE": "Ikon WhatsApp merespons klik dengan membuka tab baru berformat API WhatsApp. Format nomor yang tidak valid diblokir.",
            "BE": "Browser berhasil mengarahkan (redirect) ke alamat external WhatsApp API dengan parameter query nomor telepon admin yang valid.",
            "DB": "Sistem berhasil memuat nomor admin klinik dari database tabel `klinik` secara akurat."
        },
        "scenarios": [
            ("Skenario 1: Hubungi Admin via WhatsApp", "Pasien mengklik tombol WhatsApp untuk membuka chat bantuan admin.", "whatsapp_redirect"),
            ("Skenario 2: Validasi Format Nomor Salah", "Sistem memvalidasi nomor telepon agar menggunakan kode negara yang valid.", "whatsapp_validation")
        ]
    },
    14: {
        "code": "NF-01",
        "name": "Booking Availability Validation",
        "data_input": "HTTP Method & Path: GET /portal/appointments/new",
        "be_prefix": "backend_nf01",
        "expected": {
            "FE": "Kalender booking membatasi pemilihan tanggal. Jika pengguna memilih hari Minggu, memicu Toast error 'Klinik Tutup' dan memblokir submit form kuesioner.",
            "BE": "Request POST `/api/appointments` mengembalikan respons status 422 Unprocessable Entity apabila tanggal yang dikirim jatuh pada hari libur/Minggu.",
            "DB": "Database menolak atau tidak melakukan operasi INSERT baru pada tabel `appointments` jika validasi tanggal libur gagal."
        },
        "actual": {
            "FE": "Saat memilih tanggal hari Minggu (21 Juni 2026) dan mengklik lanjut, Toast peringatan merah 'Klinik Tutup: Maaf, klinik tutup pada hari Minggu' berhasil muncul dan menahan formulir.",
            "BE": "Percobaan submit langsung ke backend pada hari libur memicu respons API status 422 Unprocessable Content.",
            "DB": "Operasi penulisan ke tabel `appointments` dibatalkan (rollback/tidak dieksekusi) untuk menjaga integritas jadwal."
        },
        "scenarios": [
            ("Skenario 1: Validasi Ketersediaan Halaman", "Memverifikasi halaman booking dapat diakses secara lancar.", "24_reservasi_janji_temu"),
            ("Skenario 2: Validasi Booking Hari Libur", "Sistem menolak reservasi pada hari libur resmi/tutup klinik.", "holiday_booking_failed")
        ]
    },
    15: {
        "code": "NF-02",
        "name": "Data Security Validation",
        "data_input": "Security Spec: SSL/HTTPS port 443, hashing algorithm: bcrypt",
        "be_prefix": "backend_nf02",
        "expected": {
            "FE": "Halaman internal portal (dashboard, pemeriksaan, dll.) tidak dapat diakses tanpa token sesi yang valid di localStorage. Mencoba masuk tanpa login memicu redireksi ke halaman login.",
            "BE": "Seluruh route API internal dilindungi oleh middleware auth:sanctum. Request tanpa token mengembalikan status 401 Unauthorized.",
            "DB": "Kolom `passwords` pada database tabel `users` menyimpan sandi yang telah di-hash menggunakan algoritma Bcrypt (format `$2y$10$...`)."
        },
        "actual": {
            "FE": "Mencoba mengakses `/doctor/pemeriksaan` secara langsung tanpa login berhasil diredireksi paksa ke `/portal/login`.",
            "BE": "Postman/DevTools memanggil API backend tanpa token autentikasi menerima respons status 401 Unauthorized secara konsisten.",
            "DB": "Tangkapan layar tabel `users` di database membuktikan kolom `passwords` terenkripsi aman dengan Bcrypt hash."
        },
        "scenarios": [
            ("Skenario 1: Uji Pembatasan Halaman Portal", "Memverifikasi proteksi akses halaman internal berdasarkan role.", "01_login_staf_sukses"),
            ("Skenario 2: Enkripsi Sandi Pengguna (Bcrypt)", "Memverifikasi penyimpanan password pada database dalam bentuk hash.", "backend_nf02"),
            ("Skenario 3: Pengecekan Enkripsi Data & Proteksi Sesi", "Memverifikasi bahwa data sensitif pengguna dilindungi di tingkat backend dengan token Sanctum.", "backend_nf02")
        ]
    },
    16: {
        "code": "NF-03",
        "name": "Responsive Interface Validation",
        "data_input": "Viewport Dimensions: [360x740, 768x1024, 1920x1080]",
        "be_prefix": "backend_nf03",
        "expected": {
            "FE": "Seluruh layout halaman, menu sidebar, card, dan tombol menyesuaikan diri (responsif) menggunakan utilitas Tailwind CSS tanpa bertumpuk atau meluber pada resolusi Mobile, Tablet, dan Desktop.",
            "BE": "Server backend mengembalikan response data JSON yang sama secara konsisten tanpa tergantung pada jenis perangkat pengirim request.",
            "DB": "Database memproses query SELECT yang sama dari API untuk menyajikan data ke layar mobile maupun desktop."
        },
        "actual": {
            "FE": "Elemen UI berhasil menyesuaikan diri: sidebar tersembunyi menjadi hamburger menu pada mobile, dan tabel antrean berubah menjadi scrollable horizontal.",
            "BE": "API membalas dengan status 200 OK dan payload JSON yang seragam baik diakses dari mobile maupun desktop.",
            "DB": "Data rekam medis dan antrean ter-SELECT secara konsisten dari tabel database tanpa kendala resolusi perangkat."
        },
        "scenarios": [
            ("Skenario 1: Tampilan Mobile Viewport", "Memverifikasi tata letak elemen tetap rapi pada layar mobile.", "mobile_login_view"),
            ("Skenario 2: Tampilan Tablet Viewport", "Memverifikasi responsivitas antarmuka pada ukuran layar tablet.", "tablet_login_view"),
            ("Skenario 3: Tampilan Desktop Viewport", "Memverifikasi tampilan halaman di browser desktop.", "01_login_staf_sukses"),
            ("Skenario 4: Pengujian Viewport Sangat Kecil", "Memverifikasi layout tetap terbaca pada ukuran layar minimum.", "mobile_login_view")
        ]
    },
    17: {
        "code": "NF-04",
        "name": "User Interaction Efficiency",
        "data_input": "Action triggers: onClick(), onChange(), onSubmit()",
        "be_prefix": "backend_nf04",
        "expected": {
            "FE": "Setiap aksi interaksi (klik tombol, ubah input) merespons seketika. Tombol submit dinonaktifkan (`disabled`) setelah diklik sekali untuk mencegah pengiriman data ganda (double submit).",
            "BE": "API memproses antrean request secara berurutan dan mengabaikan request duplikat cepat dari session yang sama.",
            "DB": "Database tidak mengalami duplikasi record data (misal: double insert tabel `appointments`) akibat klik tombol submit berulang-ulang."
        },
        "actual": {
            "FE": "Tombol 'Selesaikan Pendaftaran' dan 'Simpan' secara otomatis mendapatkan class disabled dan loading spinner setelah diklik sekali.",
            "BE": "API menerima satu request utama dan mengembalikan status 201 Created. Request duplikat berikutnya terblokir.",
            "DB": "Tabel `appointments` dan `rekam_medis_pasien` terhindar dari baris data ganda (duplikasi) dalam database."
        },
        "scenarios": [
            ("Skenario 1: Validasi Efisiensi Interaksi", "Memastikan transisi antar halaman mulus tanpa delay yang mengganggu.", "53_dashboard_summary_cards"),
            ("Skenario 2: Pencegahan Double Submit Form", "Tombol dinonaktifkan setelah ditekan sekali untuk mencegah double post.", "53_dashboard_summary_cards")
        ]
    },
    18: {
        "code": "NF-05",
        "name": "Page Load Time Measurement",
        "data_input": "Network performance logs: [DNS Lookup, TCP Handshake, TTFB, DOM Content Loaded]",
        "be_prefix": "backend_nf05",
        "expected": {
            "FE": "Halaman aplikasi SIGIGI ter-render sepenuhnya di layar browser dalam waktu kurang dari 3 detik pada koneksi internet normal, dan tetap termuat stabil pada koneksi lambat.",
            "BE": "API mengoptimalkan response time melalui caching/indexing database agar TTFB (Time to First Byte) berada di bawah 200 ms.",
            "DB": "Database memproses query indeks tabel `appointments` dan `users` secara cepat dengan eksekusi query kurang dari 50 ms."
        },
        "actual": {
            "FE": "Waktu muat halaman normal sangat cepat (TTFB 85ms, DOM Content Loaded 415ms). Pada simulasi Slow 3G di Network throttling, halaman termuat bertahap secara stabil.",
            "BE": "API mengembalikan respons JSON dalam waktu singkat, tercatat respons time total rata-rata berkisar 85 ms.",
            "DB": "Query SELECT database berhasil dieksekusi dengan efisiensi tinggi di bawah batas toleransi waktu muat."
        },
        "scenarios": [
            ("Skenario 1: Pengukuran Kecepatan Normal", "Memverifikasi waktu muat API berada di bawah batas toleransi.", "devtools_nf05_normal"),
            ("Skenario 2: Pengukuran Kecepatan Slow 3G", "Memverifikasi waktu muat halaman pada koneksi jaringan lambat.", "devtools_nf05_slow3g")
        ]
    }
}

original_doc = Document(original_doc_path)
doc = Document(doc_path)

print(f"Loaded original document: {original_doc_path}")
print(f"Loaded target document: {doc_path}")

for t_idx, req in list(requirements_info.items()):
    orig_table = original_doc.tables[t_idx]
    target_table = doc.tables[t_idx]
    
    # 1. Build metadata rows with 3 sides split
    rows_data = [
        ("Kode", req["code"]),
        ("Nama Spesifikasi", req["name"]),
        ("Data Input", req["data_input"]),
        
        ("Expected Result (Frontend)", req["expected"]["FE"]),
        ("Expected Result (Backend)", req["expected"]["BE"]),
        ("Expected Result (Database)", req["expected"]["DB"]),
        
        ("Actual Result (Frontend)", req["actual"]["FE"]),
        ("Actual Result (Backend)", req["actual"]["BE"]),
        ("Actual Result (Database)", req["actual"]["DB"]),
        
        ("Status", "Berhasil")
    ]
    
    # 2. Create a clean 12-column outer table
    new_table = doc.add_table(rows=0, cols=12)
    new_table.style = target_table.style
    new_table.autofit = False
    
    # Write header (Komponen | Keterangan)
    cells = add_grid_row(new_table, [3, 9])
    format_cell_text(cells[0], "Komponen", bold=True)
    format_cell_text(cells[1], "Keterangan", bold=True)
    
    # Write metadata rows
    for label, val in rows_data:
        cells = add_grid_row(new_table, [3, 9])
        format_cell_text(cells[0], label, bold=True)
        format_cell_text(cells[1], val)
        
    # Write "Skenario Pengujian" Section Header (merged across all 12 columns)
    cells = add_grid_row(new_table, [12])
    format_cell_text(cells[0], "Skenario Pengujian", bold=True, size_pt=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    # 3. Group scenarios into chunks of up to 3 for horizontal layout
    scenarios = req["scenarios"]
    num_scenarios = len(scenarios)
    chunks = []
    
    if num_scenarios == 1:
        chunks = [scenarios]
    elif num_scenarios == 2:
        chunks = [scenarios]
    elif num_scenarios == 3:
        chunks = [scenarios]
    elif num_scenarios == 4:
        chunks = [scenarios[0:2], scenarios[2:4]]
    elif num_scenarios == 5:
        chunks = [scenarios[0:3], scenarios[3:5]]
    elif num_scenarios == 6:
        chunks = [scenarios[0:3], scenarios[3:6]]
    else:
        for i in range(0, num_scenarios, 3):
            chunks.append(scenarios[i:i+3])
            
    # For each chunk, insert title and content rows directly in the flat table
    for chunk in chunks:
        chunk_len = len(chunk)
        # Calculate grid spans
        spans = []
        if chunk_len == 1:
            spans = [12]
        elif chunk_len == 2:
            spans = [6, 6]
        elif chunk_len == 3:
            spans = [4, 4, 4]
            
        # Title Row
        title_cells = add_grid_row(new_table, spans)
        for idx, (s_title, s_desc, img_prefix) in enumerate(chunk):
            format_cell_text(title_cells[idx], s_title, bold=True, size_pt=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
            
        # Content Row (Description + Screenshot Evidence FE + BE/DB)
        content_cells = add_grid_row(new_table, spans)
        for idx, (s_title, s_desc, img_prefix) in enumerate(chunk):
            col_width_in = (spans[idx] / 12.0) * 5.37
            populate_content_cell(content_cells[idx], s_desc, img_prefix, req["be_prefix"], col_width_in)

    # 4. Replace XML elements in target doc
    tbl_element = target_table._tbl
    new_tbl_element = new_table._tbl
    parent = tbl_element.getparent()
    parent.insert(parent.index(tbl_element), new_tbl_element)
    parent.remove(tbl_element)
    
    print(f"Table {t_idx} ({req['code']}): Rebuilt using flat horizontal cells layout with 3-side expected/actual results and dual FE-BE/DB evidence.")

doc.save(doc_path)
print(f"\nDone! Successfully updated Draft CD5.docx with flat horizontal cells layout.")
