import sys
import copy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout.reconfigure(encoding='utf-8')

doc = Document(r'c:\College\Capstone Design\sigigi-main\Manual_Book_SIGIGI_2.0.docx')

changes = []

def replace_text_in_para(para, old, new):
    """Replace text across all runs in a paragraph."""
    full = para.text
    if old not in full:
        return False
    # Try simple single-run replace first
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    # Multi-run: rebuild in first non-empty run
    new_full = full.replace(old, new)
    first_run = None
    for run in para.runs:
        if run.text.strip():
            first_run = run
            break
    for run in para.runs:
        run.text = ''
    if first_run:
        first_run.text = new_full
    elif para.runs:
        para.runs[0].text = new_full
    return True

def add_row_to_table(table, cells_data, copy_fmt_from_row_idx=1):
    """Add a new row to a table, copying format from an existing row."""
    ref_row = table.rows[copy_fmt_from_row_idx]
    new_tr = copy.deepcopy(ref_row._tr)
    table._tbl.append(new_tr)
    new_row = table.rows[-1]
    for i, text in enumerate(cells_data):
        if i < len(new_row.cells):
            for para in new_row.cells[i].paragraphs:
                for run in para.runs:
                    run.text = ''
            if new_row.cells[i].paragraphs:
                p = new_row.cells[i].paragraphs[0]
                if p.runs:
                    p.runs[0].text = text
                else:
                    run = p.add_run(text)
            else:
                new_row.cells[i].text = text
    return new_row

# =====================================================================
# CHANGE 1: Bab 1.4 — Update catatan URL placeholder ke URL asli
# =====================================================================
for para in doc.paragraphs:
    if 'URL pada tabel di atas merupakan contoh penamaan' in para.text:
        if replace_text_in_para(para,
            'URL pada tabel di atas merupakan contoh penamaan. Pada dokumen final, sesuaikan dengan domain deployment yang digunakan.',
            'URL deployment resmi SIGIGI 2.0 adalah https://sigigi.my.id. Portal pasien dan portal admin/staff dapat diakses melalui domain yang sama dengan jalur yang sesuai.'):
            changes.append('1. Fixed 1.4 URL placeholder note')
        break

# Also update the 1.4 URL table cells if they contain placeholder text
for t_idx, table in enumerate(doc.tables[:10]):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if 'contoh' in para.text.lower() and ('url' in para.text.lower() or 'portal' in para.text.lower()):
                    if replace_text_in_para(para, para.text, para.text.replace('(contoh)', '').strip()):
                        changes.append(f'1b. Fixed URL placeholder in table {t_idx}')

# =====================================================================
# CHANGE 2: Bab 3.1 — Fix login step 1: add real URL
# =====================================================================
for para in doc.paragraphs:
    if para.text.strip() == 'Buka alamat portal staff/admin melalui browser.':
        if replace_text_in_para(para,
            'Buka alamat portal staff/admin melalui browser.',
            'Buka https://sigigi.my.id melalui browser.'):
            changes.append('2. Fixed 3.1 step 1 URL')
        break

# =====================================================================
# CHANGE 3: Bab 5.1 Registrasi — Add URL to step 1
# =====================================================================
for para in doc.paragraphs:
    if para.text.strip() == 'Buka halaman registrasi pasien.':
        if replace_text_in_para(para,
            'Buka halaman registrasi pasien.',
            'Buka https://sigigi.my.id/portal/register melalui browser.'):
            changes.append('3. Fixed 5.1 registrasi step 1 URL')
        break

# =====================================================================
# CHANGE 4: Bab 5.2 Login Pasien — Add URL to step 1
# =====================================================================
for para in doc.paragraphs:
    if para.text.strip() == 'Buka halaman login portal pasien.':
        if replace_text_in_para(para,
            'Buka halaman login portal pasien.',
            'Buka https://sigigi.my.id/portal/login melalui browser.'):
            changes.append('4. Fixed 5.2 login step 1 URL')
        break

# =====================================================================
# CHANGE 5: Bab 5.4 — Remove step "Isi keluhan singkat pada kolom keluhan utama."
# and fix the "Klik Lanjutkan" text (it's now step about choosing klinik & jadwal)
# =====================================================================
for para in doc.paragraphs:
    if para.text.strip() == 'Isi keluhan singkat pada kolom keluhan utama.':
        if replace_text_in_para(para,
            'Isi keluhan singkat pada kolom keluhan utama.',
            'Pastikan data lokasi klinik, tanggal, dan waktu kunjungan sudah benar sebelum melanjutkan.'):
            changes.append('5. Fixed 5.4 step 4 wrong keluhan field')
        break

for para in doc.paragraphs:
    if para.text.strip() == 'Klik Lanjutkan untuk mengisi questionnaire keluhan.':
        if replace_text_in_para(para,
            'Klik Lanjutkan untuk mengisi questionnaire keluhan.',
            'Klik tombol Lanjut Isi Kuesioner untuk melanjutkan ke halaman questionnaire keluhan.'):
            changes.append('5b. Fixed 5.4 step 5 button name')
        break

# =====================================================================
# CHANGE 6: Bab 5.5 Questionnaire — Fix button name & add note about upload
# =====================================================================
for para in doc.paragraphs:
    if 'Klik Lanjutkan atau Submit sesuai tombol yang tersedia.' in para.text:
        if replace_text_in_para(para,
            'Klik Lanjutkan atau Submit sesuai tombol yang tersedia.',
            'Isi deskripsi keluhan tambahan pada kolom teks (opsional).'):
            changes.append('6a. Fixed 5.5 step 4 button name')
        break

# Fix hasil yang diharapkan 5.5 to mention upload on same page
for para in doc.paragraphs:
    if para.text.strip() == 'Sistem menghitung tingkat urgensi awal seperti rendah, sedang, atau tinggi berdasarkan jawaban pasien.':
        if replace_text_in_para(para,
            'Sistem menghitung tingkat urgensi awal seperti rendah, sedang, atau tinggi berdasarkan jawaban pasien.',
            'Sistem menyimpan appointment dan menghitung tingkat urgensi awal (Rendah, Sedang, atau Tinggi) berdasarkan jawaban pasien. Apabila citra gigi diunggah, sistem akan melakukan analisis awal karies/non-karies secara otomatis.'):
            changes.append('6b. Fixed 5.5 expected result')
        break

# =====================================================================
# CHANGE 7: Bab 5.6 — Fix step 5 "Klik Submit Appointment" & merge context
# =====================================================================
for para in doc.paragraphs:
    if para.text.strip() == 'Klik Submit Appointment.':
        if replace_text_in_para(para,
            'Klik Submit Appointment.',
            'Klik tombol Selesaikan Pendaftaran untuk mengirimkan appointment.'):
            changes.append('7. Fixed 5.6 step 5 button name')
        break

# Fix 5.6 langkah 1 to reference it's on the same questionnaire page
for para in doc.paragraphs:
    if para.text.strip() == 'Siapkan citra gigi dengan format JPG, JPEG, atau PNG.':
        if replace_text_in_para(para,
            'Siapkan citra gigi dengan format JPG, JPEG, atau PNG.',
            'Pada halaman questionnaire yang sama, siapkan citra gigi dengan format JPG, JPEG, atau PNG (ukuran maksimal 10 MB).'):
            changes.append('7b. Fixed 5.6 step 1 context')
        break

# =====================================================================
# CHANGE 8: Bab 4.6 Input Tindakan — add note about tooth format
# =====================================================================
for para in doc.paragraphs:
    if para.text.strip() == 'Isi nomor gigi apabila tindakan berkaitan dengan gigi tertentu.':
        if replace_text_in_para(para,
            'Isi nomor gigi apabila tindakan berkaitan dengan gigi tertentu.',
            'Isi nomor gigi apabila tindakan berkaitan dengan gigi tertentu (contoh: 46, 36, 11). Untuk tindakan umum seperti scaling atau kontrol, kolom nomor gigi dapat dikosongkan.'):
            changes.append('8. Fixed 4.6 nomor gigi description')
        break

# =====================================================================
# CHANGE 9: Bab 3.13 — Add cross-reference note (find heading 3.13)
# =====================================================================
for i, para in enumerate(doc.paragraphs):
    if '3.13' in para.text and 'Tindakan' in para.text:
        # Find the last paragraph of this section and add a note
        # Look for next content paragraph after this heading
        for j in range(i+1, min(i+10, len(doc.paragraphs))):
            next_para = doc.paragraphs[j]
            if 'Tujuan fitur' in next_para.text:
                if replace_text_in_para(next_para,
                    next_para.text,
                    next_para.text + ' Catatan: Modul tindakan juga dapat diakses langsung melalui Portal Dokter. Lihat Bab 4 untuk panduan lengkap fitur dokter.'):
                    changes.append('9. Added cross-reference note in 3.13')
                break
        break

# =====================================================================
# CHANGE 10: Table 7.2 (Tingkat Urgensi) — Add calculation explanation
# =====================================================================
for para in doc.paragraphs:
    if 'Tingkat urgensi pada SIGIGI 2.0 dihasilkan dari jawaban questionnaire pasien' in para.text:
        if replace_text_in_para(para,
            'Tingkat urgensi pada SIGIGI 2.0 dihasilkan dari jawaban questionnaire pasien. Nilai urgensi digunakan sebagai informasi pendukung bagi dokter/staff untuk memahami kondisi awal pasien. Sistem tidak secara otomatis menggantikan keputusan prioritas klinis dokter.',
            'Tingkat urgensi pada SIGIGI 2.0 dihasilkan dari jawaban questionnaire pasien (10 pertanyaan Ya/Tidak). Perhitungan dilakukan otomatis oleh sistem berdasarkan jumlah jawaban "Ya": 0–3 jawaban Ya = Rendah, 4–7 jawaban Ya = Sedang, 8–10 jawaban Ya = Tinggi. Nilai urgensi digunakan sebagai informasi pendukung bagi dokter/staff untuk memahami kondisi awal pasien. Sistem tidak secara otomatis menggantikan keputusan prioritas klinis dokter.'):
            changes.append('10. Fixed 7.2 urgency calculation explanation')
        break

# =====================================================================
# CHANGE 11: Table 46 (urgency detail table) — Add calculation column hint
# =====================================================================
for t_idx, table in enumerate(doc.tables):
    # Find urgency table by checking for Rendah/Sedang/Tinggi rows
    if len(table.rows) >= 4:
        texts = [row.cells[0].text.strip() for row in table.rows]
        if 'Rendah' in texts and 'Sedang' in texts and 'Tinggi' in texts:
            for row in table.rows:
                if row.cells[0].text.strip() == 'Rendah':
                    for para in row.cells[1].paragraphs:
                        if replace_text_in_para(para,
                            'Keluhan relatif ringan berdasarkan jawaban pasien.',
                            'Keluhan relatif ringan (0–3 jawaban Ya dari 10 pertanyaan).'):
                            changes.append('11a. Fixed urgency table Rendah')
                if row.cells[0].text.strip() == 'Sedang':
                    for para in row.cells[1].paragraphs:
                        if replace_text_in_para(para,
                            'Terdapat beberapa keluhan yang perlu diperhatikan.',
                            'Terdapat beberapa keluhan yang perlu diperhatikan (4–7 jawaban Ya dari 10 pertanyaan).'):
                            changes.append('11b. Fixed urgency table Sedang')
                if row.cells[0].text.strip() == 'Tinggi':
                    for para in row.cells[1].paragraphs:
                        if replace_text_in_para(para,
                            'Keluhan menunjukkan kondisi yang berpotensi membutuhkan perhatian lebih cepat.',
                            'Keluhan menunjukkan kondisi yang berpotensi membutuhkan perhatian lebih cepat (8–10 jawaban Ya dari 10 pertanyaan).'):
                            changes.append('11c. Fixed urgency table Tinggi')
            break

# =====================================================================
# CHANGE 12: Table 48 (Troubleshooting) — Add "slot penuh" row
# =====================================================================
for t_idx, table in enumerate(doc.tables):
    # Find troubleshooting table
    if table.rows and table.rows[0].cells[0].text.strip() == 'Masalah':
        # Check if slot penuh already exists
        has_slot = any('slot' in row.cells[0].text.lower() or 'jam' in row.cells[0].text.lower() 
                       for row in table.rows[1:])
        if not has_slot:
            add_row_to_table(table, [
                'Jam kunjungan tidak tersedia atau slot penuh',
                'Jam yang dipilih sudah dipesan pasien lain untuk tanggal yang sama.',
                'Pilih jam kunjungan lain yang masih tersedia pada halaman form appointment.'
            ], copy_fmt_from_row_idx=1)
            changes.append('12. Added slot penuh troubleshooting row')
        break

# =====================================================================
# CHANGE 13: Table 42 (Alur Pasien) — Merge step 3 & 4 (questionnaire + upload)
# =====================================================================
for t_idx, table in enumerate(doc.tables):
    if table.rows and table.rows[0].cells[0].text.strip() == 'Tahap' and len(table.rows) >= 8:
        # Check if this is the alur pasien table
        if any('Pasien' in row.cells[1].text for row in table.rows[1:] if len(row.cells) > 1):
            for row in table.rows:
                if row.cells[0].text.strip() == '3':
                    for para in row.cells[2].paragraphs:
                        if 'questionnaire' in para.text.lower() or 'keluhan' in para.text.lower():
                            replace_text_in_para(para,
                                para.text,
                                'Mengisi questionnaire keluhan dan mengunggah citra gigi (opsional).')
                            changes.append('13a. Fixed alur pasien step 3 activity')
                    for para in row.cells[3].paragraphs:
                        replace_text_in_para(para,
                            para.text,
                            'Sistem menghasilkan tingkat urgensi awal. Apabila foto diunggah, sistem melakukan analisis awal karies/non-karies.')
                        changes.append('13b. Fixed alur pasien step 3 output')
                if row.cells[0].text.strip() == '4':
                    for para in row.cells[1].paragraphs:
                        if 'Pasien' in para.text:
                            replace_text_in_para(para, para.text, 'Pasien')
                    for para in row.cells[2].paragraphs:
                        if 'citra' in para.text.lower() or 'foto' in para.text.lower() or 'unggah' in para.text.lower():
                            replace_text_in_para(para,
                                para.text,
                                'Melihat antrean dan status kunjungan.')
                            changes.append('13c. Merged alur pasien step 4 (was upload, now antrean)')
                    for para in row.cells[3].paragraphs:
                        if 'ml api' in para.text.lower() or 'karies' in para.text.lower():
                            replace_text_in_para(para,
                                para.text,
                                'Nomor antrean dan status kunjungan terlihat.')
                            changes.append('13d. Fixed alur pasien step 4 output')
                # Renumber subsequent steps accordingly
                if row.cells[0].text.strip() == '5':
                    for para in row.cells[2].paragraphs:
                        if 'antrean' in para.text.lower():
                            replace_text_in_para(para,
                                para.text,
                                'Dokter memeriksa pasien dan menginput tindakan, resep, odontogram, dan catatan dokter.')
                            changes.append('13e. Fixed alur pasien step 5')
                    for para in row.cells[3].paragraphs:
                        if 'antrean' in para.text.lower() or 'nomor' in para.text.lower():
                            replace_text_in_para(para,
                                para.text,
                                'Rekam medis terbentuk.')
                            changes.append('13f. Fixed alur pasien step 5 output')
            break

# =====================================================================
# CHANGE 14: Fix Checklist 9.1 — Update with real URL
# =====================================================================
for para in doc.paragraphs:
    if '[ ] Semua URL contoh sudah diganti dengan URL deployment asli.' in para.text:
        if replace_text_in_para(para,
            '[ ] Semua URL contoh sudah diganti dengan URL deployment asli.',
            '[✔] URL deployment sudah diperbarui menjadi https://sigigi.my.id.'):
            changes.append('14. Updated checklist 9.1 URL item')
        break

# =====================================================================
# CHANGE 15: Fix 8 Troubleshooting - add note about max upload size
# =====================================================================
for t_idx, table in enumerate(doc.tables):
    if table.rows and table.rows[0].cells[0].text.strip() == 'Masalah':
        for row in table.rows:
            if 'Upload citra gagal' in row.cells[0].text:
                for para in row.cells[2].paragraphs:
                    if 'JPG' in para.text:
                        replace_text_in_para(para,
                            'Gunakan JPG/JPEG/PNG, periksa koneksi, dan ulangi upload.',
                            'Gunakan format JPG/JPEG/PNG dengan ukuran maksimal 10 MB. Periksa koneksi internet dan ulangi upload. Foto dari kamera HP berukuran besar sebaiknya dikompres terlebih dahulu.')
                        changes.append('15. Updated upload troubleshooting with size limit info')
                break
        break

# =====================================================================
# Save output
# =====================================================================
output_path = r'c:\College\Capstone Design\sigigi-main\Manual_Book_SIGIGI_2.0_REVISED.docx'
doc.save(output_path)

print(f'SAVED: {output_path}')
print(f'\nTotal changes made: {len(changes)}')
for c in changes:
    print(f'  - {c}')
